#!/usr/bin/env python
# coding: utf-8

# In[12]:


import openpyxl
wb = openpyxl.load_workbook('WorkSpace.xlsx') #输入excel文件名字
sheets = wb.sheetnames


# In[13]:


WordsList = []
for i in range(len(sheets)):
    sheet = wb[sheets[i]]
    words = [sheet.cell(row = i, column = 1).value for i in range(1,sheet.max_row + 1)]
    WordsList.append(words)


# In[14]:


import docx
import random
file=docx.Document()
file.add_heading("小测",level=1)
quiz = ""
for i in range(len(sheets)):
    if int(sheets[i][2]) == 3:
        for j in range(7):
            file.add_heading(sheets[i],level=2)
            day = "第" + str(j+1) + "天"
            file.add_heading(day, level=3)
            random.shuffle(WordsList[i])
            if len(WordsList[i]) > 50:
                quiz = "\n".join(WordsList[i][:50])
            else:
                quiz ="\n".join(WordsList[i])
            file.add_paragraph(quiz)
            file.add_page_break()
    else:
        file.add_heading(sheets[i],level=2)
        random.shuffle(WordsList[i])
        if len(WordsList[i]) > 50:
            quiz = "\n".join(WordsList[i][:50])
        else:
            quiz ="\n".join(WordsList[i])
        file.add_paragraph(quiz)
        file.add_page_break()
file.save("小测.docx") #输出的word名字


# In[ ]:




