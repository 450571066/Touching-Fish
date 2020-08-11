#!/usr/bin/env python
# coding: utf-8

# In[20]:


import requests
import pandas as pd
import re
from bs4 import BeautifulSoup
import docx
import time


# # 获取URL

# In[21]:


url = 'https://www.xbiquge.cc/book/49549/35899003.html' #测试用 章节
#main_url =  'https://www.xbiquge.cc/book/49549/'       #豪婿
main_url = 'https://www.xbiquge.cc/book/420/'           #龙王传说
header = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/65.0.3314.0 Safari/537.36 SE 2.X MetaSr 1.0'
}


# In[33]:


t = time.time()
while True:
    try:
        #加入timeout防止单次访问用时过长，timeout = 5防止对服务器造成过大负担
        main_response = requests.get(main_url, headers = header, timeout = 5)
        break
    except:
        print('---URL {} Failed...Try to connect again'.format(main_url))
print('---------Done!: used {:.2f}s--------'.format(time.time() - t))


# # 解码 乱码->中文

# In[23]:


#www.xbiquge.cc 使用gbk进行编码，所以需要gbk解码
main_response.encoding = "gbk"
main_text = main_response.content.decode('gbk')


# # 读入beautifulsoup进行解析

# In[24]:


#html.parser是解析器，也可是lxml
main_soup = BeautifulSoup(main_text,'lxml')
#把整个网页转换成解析过的String，不过感觉不需要这一步
#直接使用main_text可以达到同样效果
strMain = main_soup.prettify()
#获得主标题
find_title = re.compile(r'<meta content="(.*?)" property="og:title"/>')
#获得作者
find_author = re.compile(r'<meta content="(.*?)" property="og:novel:author"/>')
mainTitle = find_title.findall(strMain)[0]
author = find_author.findall(strMain)[0]


# # 获取小说list

# In[25]:


#解析主网页的小说列表，笔趣阁都是以同一个形式加载的
pattern = re.compile(r'<a href="(\d+.html)">')
text_list = pattern.findall(main_soup.prettify())


# In[30]:


#test block


# # 每一个章节分别进行提取输出操作

#    ## 1. 读取每一个章节的内容和标题

# In[ ]:


def write_to_doc(book):
    begin_time = time.time()
    chapter_url = main_url+book
    print('---Getting URL...---')
    while True:
        try:
            #加timeout防单词用时过长
            chapter_response = requests.get(chapter_url,headers = header,timeout = 5)
            break
        except:
            print('---URL {} Failed...Try to connect again'.format(book))
            continue
    print('-----URL {} used {:.2f}s-----'.format(book, time.time() - begin_time))
    #小说主体部分需要重新解码
    chapter_response.encoding = "gbk"
    chapter_text = chapter_response.content.decode('gbk')
    
    #以beautiful的形式读入，方便后续解析
    chapter_soup = BeautifulSoup(chapter_text,'lxml')

    #通过分析页面，笔趣阁小说主体部分在 id为content的div里
    div = chapter_soup.find_all('div',id='content')#,string='更多'
    #将获得的tag们全部变成string
    div_text =div[0].prettify()
    #网页中每个段落之间会有<br/>，刚好适合用来分段落
    content = div_text.split('<br/>')[1:]
    #再将过多的\n去除
    content = "".join(content).split('\n')[:-2]
    
    #获得章节名、文章名、以及来源
    title = chapter_soup.title.text
    mainTitle = title.split('_')[1]
    source = title.split('_')[2]
    chapterTitle = title.split('_')[0]
    
    print('---Fetched the contents---')
    
    return content, chapterTitle


# ## 2.写入word文档

# In[ ]:


begin = time.time()
#新建一个docx文件
file=docx.Document()
#写入主标题和作者名
file.add_heading("{}".format(mainTitle),level=1)
file.add_heading("作者:{}".format(author),level=2)
count = 0
for i in text_list:
    #count 用来计数，数字为章节数
    count  += 1
    print('-------------Doing chapter: {}------------'.format(count))
    
    content, chapterTitle = write_to_doc(i)
    #写入段落标题
    file.add_heading("{}".format(chapterTitle),level=2)
    
    print('----begin to write chapter {} into doc----'.format(count))
    #将爬取回的内容一段一段写入docx
    for i in content:
        if i !=  ' ':
            file.add_paragraph(i)
            
    print('---Saving for chapter {}!---'.format(count))
    #加入分页，每一章的开头都是全新的一页纸
    file.add_page_break()
    #保存docx文件
    file.save("{}.docx".format(mainTitle)) #输出的word名字
    #sleep(1)防止对服务器造成过大负担
    time.sleep(1)
    print('---Done for chapter: {}---\n'.format(count))
print('--------------------------------DONE!!!!!!!-------------------------')
print('------------------uesd: {:.2f}s-------------------'.format(time.time() - begin))


# In[ ]:




